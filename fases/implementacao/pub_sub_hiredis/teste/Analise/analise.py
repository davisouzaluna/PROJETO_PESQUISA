import redis
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import skew, kurtosis
import csv
from docx import Document
from docx.shared import Inches
import os

# Nome da chave Redis
redis_key = 'quic'

# Criar diretório para armazenar os arquivos
output_dir = f'{redis_key}'
os.makedirs(output_dir, exist_ok=True)

# Conectar ao Redis
r = redis.Redis(host='localhost', port=6379, db=0)

# Recuperar os dados da chave Redis
latencies = r.lrange(redis_key, 0, -1)
latencies = [float(latency.decode('utf-8').strip('"')) for latency in latencies]  # Convertendo para floats

# Análises Estatísticas
mean_latency = np.mean(latencies)
median_latency = np.median(latencies)
std_latency = np.std(latencies)
p95_latency = np.percentile(latencies, 95)
p99_latency = np.percentile(latencies, 99)
max_latency = np.max(latencies)
min_latency = np.min(latencies)
variance_latency = np.var(latencies)
skewness_latency = skew(latencies)
kurtosis_latency = kurtosis(latencies)

# Gerar e Salvar Gráficos
plt.style.use('seaborn-darkgrid')  # Usar estilo do seaborn para gráficos bonitos

# Histograma
plt.figure(figsize=(12, 8))
plt.hist(latencies, bins=50, alpha=0.75, color='royalblue', edgecolor='black')
plt.title('Histograma das Latências', fontsize=16)
plt.xlabel('Latência (ms)', fontsize=14)
plt.ylabel('Frequência', fontsize=14)
plt.grid(True)
plt.tight_layout()
histogram_path = os.path.join(output_dir, f'histograma_{redis_key}.png')
plt.savefig(histogram_path, dpi=300)  # Salvar com resolução maior
plt.close()

# Boxplot
plt.figure(figsize=(12, 8))
plt.boxplot(latencies, patch_artist=True, 
            boxprops=dict(facecolor='lightcyan', color='blue'),
            whiskerprops=dict(color='blue'),
            capprops=dict(color='blue'),
            medianprops=dict(color='red'))
plt.title('Boxplot das Latências', fontsize=16)
plt.ylabel('Latência (ms)', fontsize=14)
plt.grid(True)
plt.tight_layout()
boxplot_path = os.path.join(output_dir, f'boxplot_{redis_key}.png')
plt.savefig(boxplot_path, dpi=300)  # Salvar com resolução maior
plt.close()

# Gráfico da Média ao Longo do Tempo
plt.figure(figsize=(12, 8))
interval = 10  # Ajuste conforme necessário
mean_latencies = [np.mean(latencies[i:i + interval]) for i in range(0, len(latencies), interval)]
intervals = range(len(mean_latencies))  # Intervalos no eixo x
plt.plot(intervals, mean_latencies, marker='o', linestyle='-', color='coral', markersize=6, linewidth=2)
plt.title('Média das Latências ao Longo do Tempo', fontsize=16)
plt.xlabel('Intervalos', fontsize=14)
plt.ylabel('Média da Latência (ms)', fontsize=14)
plt.grid(True)
plt.tight_layout()
mean_time_path = os.path.join(output_dir, f'media_{redis_key}_tempo.png')
plt.savefig(mean_time_path, dpi=300)  # Salvar com resolução maior
plt.close()

# Salvar todos os gráficos em um único arquivo de imagem
complete_analysis_path = os.path.join(output_dir, f'analise_completa_{redis_key}.png')
fig, axs = plt.subplots(2, 2, figsize=(16, 12), constrained_layout=True)

# Histograma
img = plt.imread(histogram_path)
axs[0, 0].imshow(img)
axs[0, 0].axis('off')
axs[0, 0].set_title('Histograma das Latências', fontsize=16)

# Boxplot
img = plt.imread(boxplot_path)
axs[0, 1].imshow(img)
axs[0, 1].axis('off')
axs[0, 1].set_title('Boxplot das Latências', fontsize=16)

# Média ao Longo do Tempo
img = plt.imread(mean_time_path)
axs[1, 0].imshow(img)
axs[1, 0].axis('off')
axs[1, 0].set_title('Média das Latências ao Longo do Tempo', fontsize=16)

# Espaço vazio no canto inferior direito
axs[1, 1].axis('off')

plt.savefig(complete_analysis_path, dpi=300, bbox_inches='tight')  # Ajustar para usar todo o espaço disponível
plt.close()

# Exportar Dados em um Arquivo CSV
csv_path = os.path.join(output_dir, f'latencias_{redis_key}.csv')
with open(csv_path, 'w', newline='') as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(['Latência (ms)'])
    for latency in latencies:
        writer.writerow([latency])

# Criar Documento .docx e Inserir Análises e Imagens
doc = Document()
doc.add_heading(f'Análise de Latências - {redis_key}', level=1)

# Adicionar Análises Estatísticas
doc.add_heading('Resultados Estatísticos', level=2)
doc.add_paragraph(f"Média das Latências: {mean_latency:.6f} ms")
doc.add_paragraph("A média é o valor médio das latências registradas. Ela fornece uma visão geral do tempo médio de resposta.")
doc.add_paragraph(f"Mediana das Latências: {median_latency:.6f} ms")
doc.add_paragraph("A mediana é o valor central das latências, dividindo os dados em duas partes iguais. É útil para entender a latência típica sem ser influenciado por valores extremos.")
doc.add_paragraph(f"Desvio Padrão: {std_latency:.6f} ms")
doc.add_paragraph("O desvio padrão mede a dispersão das latências em torno da média. Valores altos indicam maior variabilidade nas latências.")
doc.add_paragraph(f"Percentil 95: {p95_latency:.6f} ms")
doc.add_paragraph("O percentil 95 indica que 95% das latências estão abaixo desse valor. É uma métrica importante para entender a latência máxima aceitável para a maioria das solicitações.")
doc.add_paragraph(f"Percentil 99: {p99_latency:.6f} ms")
doc.add_paragraph("O percentil 99 mostra que 99% das latências estão abaixo desse valor. Ajuda a identificar latências extremamente altas que ocorrem raramente.")
doc.add_paragraph(f"Latência Máxima: {max_latency:.6f} ms")
doc.add_paragraph("A latência máxima é o maior valor registrado de latência. É importante para identificar os piores cenários de desempenho.")
doc.add_paragraph(f"Latência Mínima: {min_latency:.6f} ms")
doc.add_paragraph("A latência mínima é o menor valor registrado. Mostra o melhor desempenho registrado.")
doc.add_paragraph(f"Variância: {variance_latency:.6f} ms²")
doc.add_paragraph("A variância é a medida da dispersão das latências em torno da média, elevada ao quadrado. Ela fornece uma ideia de como os dados estão espalhados.")
doc.add_paragraph(f"Assimetria: {skewness_latency:.6f}")
doc.add_paragraph("A assimetria (skewness) indica a simetria da distribuição das latências. Valores positivos sugerem que a distribuição tem cauda à direita, e valores negativos sugerem cauda à esquerda.")
doc.add_paragraph(f"Curtose: {kurtosis_latency:.6f}")
doc.add_paragraph("A curtose mede a 'altitude' da distribuição das latências. Valores positivos indicam caudas mais pesadas do que uma distribuição normal, e valores negativos indicam caudas mais leves.")

# Adicionar Gráficos ao Documento
doc.add_heading('Gráficos', level=2)

doc.add_heading('Histograma das Latências', level=3)
doc.add_picture(histogram_path, width=Inches(6))
doc.add_paragraph("O histograma mostra a distribuição das latências. É útil para visualizar a frequência de diferentes intervalos de latência.")

doc.add_heading('Boxplot das Latências', level=3)
doc.add_picture(boxplot_path, width=Inches(6))
doc.add_paragraph("O boxplot fornece uma visão geral da distribuição das latências, incluindo mediana, quartis e possíveis outliers.")

doc.add_heading('Média das Latências ao Longo do Tempo', level=3)
doc.add_picture(mean_time_path, width=Inches(6))
doc.add_paragraph("O gráfico da média ao longo do tempo mostra como a média das latências muda em diferentes intervalos de tempo.")

doc.add_heading('Análise Completa', level=3)
doc.add_picture(complete_analysis_path, width=Inches(6))
doc.add_paragraph("A análise completa inclui todos os gráficos em um único arquivo para uma visão geral abrangente.")

# Salvar Documento
doc.save(os.path.join(output_dir, f'analise_completa_{redis_key}.docx'))
